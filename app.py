import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import io
import os
from datetime import datetime
from PIL import Image
import tempfile

# ==========================================================
# 1. CONFIGURACIÓN DE LA PÁGINA STREAMLIT
# ==========================================================
st.set_page_config(
    page_title="INT – Laboratorio de Función Pulmonar",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Estilo CSS personalizado
st.markdown("""
    <style>
    .main-header {
        text-align: center;
        color: #1E3A8A;
        margin-bottom: 2rem;
    }
    .stButton>button {
        background-color: #1E3A8A;
        color: white;
        font-weight: bold;
        width: 100%;
        padding: 0.75rem;
    }
    .success-box {
        background-color: #D1FAE5;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #10B981;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #FEF3C7;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #F59E0B;
        margin: 1rem 0;
    }
    .metric-box {
        background-color: #F3F4F6;
        padding: 0.5rem;
        border-radius: 0.25rem;
        text-align: center;
        margin: 0.25rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================================
# 2. EXTRACCIÓN DE DATOS DEL PDF SUNVOU (MEJORADA Y SEGURA)
# ==========================================================
def extraer_datos_pdf_sunvou(pdf_file):
    """
    Extrae datos y gráficos del informe Sunvou con mayor precisión
    y manejo seguro de valores nulos
    """
    try:
        pdf_bytes = pdf_file.read()
        doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        pagina = doc_pdf[0]
        texto_completo = pagina.get_text()
        
        # Mostrar texto extraído para debugging (opcional)
        if st.session_state.get('debug_mode', False):
            with st.expander("📄 Texto extraído del PDF"):
                st.text(texto_completo[:2000])  # Primeros 2000 caracteres
        
        # Patrones mejorados para extracción
        def buscar_valor(patron, texto=texto_completo, default="---"):
            try:
                matches = re.findall(patron, texto, re.IGNORECASE | re.MULTILINE | re.DOTALL)
                if matches:
                    # Tomar el último match (el más probable)
                    valor = str(matches[-1]).strip()
                    # Limpiar el valor
                    valor = re.sub(r'[^\d.,]', '', valor)
                    return valor if valor else default
                return default
            except:
                return default
        
        # Extracción de valores principales con múltiples patrones
        datos = {
            "FeNO50": "---",
            "Temperatura": "---",
            "Presion": "---", 
            "Flujo": "---",
            "img_curva": None,
            "img_logo": None
        }
        
        # Múltiples patrones para cada valor (para cubrir diferentes formatos)
        patrones_feno = [
            r'FeN[O0]50[:\s]*(\d+[\.,]?\d*)',
            r'FeNO50[:\s]*(\d+[\.,]?\d*)',
            r'FeN050[:\s]*(\d+[\.,]?\d*)',
            r'FeNO\s*50[:\s]*(\d+[\.,]?\d*)',
            r'FeNO50\s*(\d+[\.,]?\d*)'
        ]
        
        patrones_temp = [
            r'Temperatura[:\s]*(\d+[\.,]?\d*)',
            r'Temp\.?[:\s]*(\d+[\.,]?\d*)',
            r'T[:\s]*(\d+[\.,]?\d*)\s*°C'
        ]
        
        patrones_pres = [
            r'Presión[:\s]*(\d+[\.,]?\d*)',
            r'Pres\.?[:\s]*(\d+[\.,]?\d*)',
            r'Pres[:\s]*(\d+[\.,]?\d*)\s*cm'
        ]
        
        patrones_flujo = [
            r'Tasa de Flujo[:\s]*(\d+[\.,]?\d*)',
            r'Tasa de flujo[:\s]*(\d+[\.,]?\d*)',
            r'Flujo[:\s]*(\d+[\.,]?\d*)',
            r'Flow[:\s]*(\d+[\.,]?\d*)'
        ]
        
        # Función para buscar con múltiples patrones
        def buscar_con_patrones(patrones):
            for patron in patrones:
                valor = buscar_valor(patron)
                if valor != "---":
                    return valor
            return "---"
        
        # Buscar cada valor
        datos["FeNO50"] = buscar_con_patrones(patrones_feno)
        datos["Temperatura"] = buscar_con_patrones(patrones_temp)
        datos["Presion"] = buscar_con_patrones(patrones_pres)
        datos["Flujo"] = buscar_con_patrones(patrones_flujo)
        
        # Limpiar y formatear valores de forma segura
        def limpiar_valor(valor):
            if valor == "---" or valor is None:
                return "---"
            try:
                # Reemplazar coma por punto para decimales
                valor_str = str(valor).replace(',', '.')
                # Eliminar caracteres no numéricos excepto punto decimal
                valor_str = re.sub(r'[^\d.]', '', valor_str)
                # Verificar si es un número válido
                if valor_str and valor_str != '.':
                    # Si no tiene decimales, agregar .0
                    if '.' not in valor_str:
                        return f"{valor_str}"
                    return valor_str
                return "---"
            except:
                return "---"
        
        # Aplicar limpieza
        for key in ['FeNO50', 'Temperatura', 'Presion', 'Flujo']:
            datos[key] = limpiar_valor(datos[key])
        
        # Buscar y extraer logos si existen
        try:
            imagenes = pagina.get_images(full=True)
            for img_index, img_info in enumerate(imagenes[:5]):  # Limitar a 5 imágenes
                try:
                    xref = img_info[0]
                    img = doc_pdf.extract_image(xref)
                    img_bytes = img["image"]
                    
                    # Filtrar por tamaño para identificar logos
                    if img["width"] > 100 and img["height"] > 30:
                        datos["img_logo"] = img_bytes
                        break  # Tomar solo el primer logo
                    
                except Exception as e:
                    continue
        except:
            datos["img_logo"] = None
        
        # Extraer específicamente la curva de exhalación
        datos["img_curva"] = extraer_curva_exhalacion_mejorada(pagina, texto_completo)
        
        doc_pdf.close()
        
        # Log de depuración
        st.session_state['ultimos_datos'] = datos
        
        return datos
        
    except Exception as e:
        st.error(f"Error procesando PDF: {str(e)}")
        # Devolver datos por defecto para evitar errores
        return {
            "FeNO50": "---",
            "Temperatura": "---",
            "Presion": "---",
            "Flujo": "---",
            "img_curva": None,
            "img_logo": None
        }

def extraer_curva_exhalacion_mejorada(pagina, texto=""):
    """
    Extrae la curva de exhalación con coordenadas precisas
    """
    try:
        # Buscar coordenadas basadas en texto
        bloques = pagina.get_text("blocks")
        
        # Buscar texto relacionado con la curva
        for bloque in bloques:
            x0, y0, x1, y1, texto_bloque, *_ = bloque
            
            if any(keyword in texto_bloque.upper() for keyword in ["CURVA", "EXHALACIÓN", "EXHALACION", "GRAPH", "PLOT"]):
                # Ajustar coordenadas para capturar la curva
                # Estas coordenadas son aproximadas, pueden necesitar ajuste
                rect_curva = fitz.Rect(
                    x0 - 20,      # Margen izquierdo
                    y1 + 10,      # Bajar desde el texto
                    min(x0 + 250, pagina.rect.width),  # Ancho máximo
                    min(y1 + 150, pagina.rect.height)  # Alto máximo
                )
                
                # Capturar la imagen de la curva
                pix = pagina.get_pixmap(
                    clip=rect_curva,
                    matrix=fitz.Matrix(2, 2),  # Resolución media
                    alpha=False
                )
                
                # Convertir a bytes PNG
                img_bytes = pix.tobytes("png")
                
                # Opcional: procesar con PIL para ajustes
                try:
                    img_pil = Image.open(io.BytesIO(img_bytes))
                    # Recortar bordes blancos si es necesario
                    # (opcional, descomentar si es necesario)
                    # img_pil = recortar_bordes(img_pil)
                    
                    output = io.BytesIO()
                    img_pil.save(output, format='PNG')
                    return output.getvalue()
                except:
                    return img_bytes
                
                break
        
        # Si no encontró por texto, usar área predeterminada
        rect_curva = fitz.Rect(50, 300, 400, 500)
        pix = pagina.get_pixmap(
            clip=rect_curva,
            matrix=fitz.Matrix(2, 2),
            alpha=False
        )
        return pix.tobytes("png")
        
    except Exception as e:
        st.warning(f"No se pudo extraer la curva: {str(e)}")
        return None

# Función auxiliar para recortar bordes blancos (opcional)
def recortar_bordes(imagen):
    """Recorta bordes blancos de una imagen"""
    try:
        from PIL import ImageChops
        
        bg = Image.new(imagen.mode, imagen.size, (255, 255, 255))
        diff = ImageChops.difference(imagen, bg)
        diff = ImageChops.add(diff, diff, 2.0, -100)
        bbox = diff.getbbox()
        
        if bbox:
            return imagen.crop(bbox)
        return imagen
    except:
        return imagen

# ==========================================================
# 3. GENERACIÓN DEL DOCUMENTO WORD (IDÉNTICO AL MODELO)
# ==========================================================
def generar_documento_identico(datos, output_path="informe_generado.docx"):
    """
    Crea un documento Word idéntico al modelo proporcionado
    con manejo seguro de datos
    """
    try:
        # Crear nuevo documento
        doc = Document()
        
        # ==================== CONFIGURACIÓN DEL DOCUMENTO ====================
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # ==================== CABECERA CON LOGOS ====================
        # Primera fila de logos
        header_table1 = doc.add_table(rows=1, cols=2)
        header_table1.autofit = False
        header_table1.columns[0].width = Inches(3.5)
        header_table1.columns[1].width = Inches(4.5)
        
        cell_left = header_table1.cell(0, 0)
        cell_right = header_table1.cell(0, 1)
        
        # Logo izquierdo (si existe)
        if datos.get("img_logo"):
            try:
                paragraph_left = cell_left.paragraphs[0]
                run_left = paragraph_left.add_run()
                run_left.add_picture(io.BytesIO(datos["img_logo"]), width=Inches(2.2))
            except:
                # Si falla, dejar celda vacía
                pass
        
        # Logo derecho (si existe)
        paragraph_right = cell_right.paragraphs[0]
        paragraph_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if datos.get("img_logo"):
            try:
                run_right = paragraph_right.add_run()
                run_right.add_picture(io.BytesIO(datos["img_logo"]), width=Inches(2.9))
            except:
                pass
        
        doc.add_paragraph()  # Espacio
        
        # ==================== TABLA DE DATOS DEL PACIENTE ====================
        datos_table = doc.add_table(rows=8, cols=4)
        datos_table.style = 'Table Grid'
        datos_table.autofit = False
        
        # Configurar anchos de columnas
        for i, width in enumerate([1.5, 3.0, 1.5, 3.0]):
            datos_table.columns[i].width = Inches(width)
        
        # Función segura para obtener datos
        def get_dato(key, default="---"):
            valor = datos.get(key, default)
            return str(valor) if valor is not None else default
        
        # Llenar datos del paciente
        filas_datos = [
            ("Nombre:", get_dato("nombre"), "Apellidos:", get_dato("apellidos")),
            ("RUT:", get_dato("rut"), "Género:", get_dato("genero")),
            ("Operador:", get_dato("operador"), "Médico:", get_dato("medico")),
            ("F. nacimiento:", get_dato("f_nacimiento"), "Edad:", get_dato("edad")),
            ("Altura:", f"{get_dato('altura')} cm", "Peso:", f"{get_dato('peso')} kg"),
            ("Raza:", "Caucásica", "Procedencia:", get_dato("procedencia")),
            ("", "", "", ""),
            ("Fecha de Examen:", get_dato("fecha_examen"), "", "")
        ]
        
        for i, (cell1, cell2, cell3, cell4) in enumerate(filas_datos):
            datos_table.cell(i, 0).text = cell1
            datos_table.cell(i, 1).text = cell2
            datos_table.cell(i, 2).text = cell3
            datos_table.cell(i, 3).text = cell4
            
            # Aplicar formato a las celdas
            for j in range(4):
                cell = datos_table.cell(i, j)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
        
        doc.add_paragraph()  # Espacio
        
        # ==================== TÍTULO PRUEBA ====================
        titulo = doc.add_paragraph()
        titulo_run = titulo.add_run("Prueba de Óxido Nítrico Exhalado")
        titulo_run.bold = True
        titulo_run.font.size = Pt(14)
        titulo_run.font.name = 'Calibri'
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ==================== INFORMACIÓN DEL EQUIPO ====================
        equipo = doc.add_paragraph()
        equipo_run = equipo.add_run("Predictivos: ATS/ERS Equipo: CA2122 FeNO (Sunvou)")
        equipo_run.font.size = Pt(11)
        equipo_run.font.name = 'Calibri'
        equipo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio
        
        # ==================== TABLA DE PARÁMETROS TÉCNICOS ====================
        params_table = doc.add_table(rows=3, cols=3)
        params_table.style = 'Table Grid'
        
        # Configurar anchos
        params_table.columns[0].width = Inches(2.0)
        params_table.columns[1].width = Inches(2.0)
        params_table.columns[2].width = Inches(1.5)
        
        parametros = [
            ("Temperatura:", get_dato("temperatura"), "°C"),
            ("Presión:", get_dato("presion"), "cmH2O"),
            ("Tasa de Flujo:", get_dato("flujo"), "ml/s")
        ]
        
        for i, (param, valor, unidad) in enumerate(parametros):
            params_table.cell(i, 0).text = param
            params_table.cell(i, 1).text = valor
            params_table.cell(i, 2).text = unidad
            
            # Formato de celdas
            for j in range(3):
                cell = params_table.cell(i, j)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
        
        doc.add_paragraph()  # Espacio
        
        # ==================== RESULTADO FeNO50 ====================
        # Tabla para FeNO50
        feno_table = doc.add_table(rows=1, cols=2)
        feno_table.autofit = False
        feno_table.columns[0].width = Inches(2.0)
        feno_table.columns[1].width = Inches(2.0)
        
        feno_cell1 = feno_table.cell(0, 0)
        feno_cell2 = feno_table.cell(0, 1)
        
        # Título FeNO50 con superíndice
        p1 = feno_cell1.paragraphs[0]
        run1 = p1.add_run("FeNO")
        run1.font.name = 'Calibri'
        run1.font.size = Pt(12)
        run1.bold = True
        
        # Añadir superíndice para el 50
        run_super = p1.add_run("50")
        run_super.font.name = 'Calibri'
        run_super.font.size = Pt(9)
        run_super.font.superscript = True
        
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Valor FeNO50
        p2 = feno_cell2.paragraphs[0]
        run2 = p2.add_run(get_dato("feno50"))
        run2.font.name = 'Calibri'
        run2.font.size = Pt(16)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0, 0, 0)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio
        
        # ==================== CURVA DE EXHALACIÓN ====================
        # Título de la curva
        titulo_curva = doc.add_paragraph()
        titulo_curva_run = titulo_curva.add_run("Curva de Exhalación FeNO50:")
        titulo_curva_run.bold = True
        titulo_curva_run.font.name = 'Calibri'
        titulo_curva_run.font.size = Pt(12)
        titulo_curva.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Insertar la curva si existe
        if datos.get("img_curva"):
            try:
                # Crear tabla para centrar la imagen
                img_table = doc.add_table(rows=1, cols=1)
                img_table.autofit = False
                img_table.columns[0].width = Inches(6.0)
                
                cell_img = img_table.cell(0, 0)
                p_img = cell_img.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run_img = p_img.add_run()
                # Ajustar tamaño según el modelo
                run_img.add_picture(
                    io.BytesIO(datos["img_curva"]), 
                    width=Inches(3.7)
                )
                
            except Exception as e:
                # Insertar placeholder si hay error
                p_placeholder = doc.add_paragraph()
                p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
                placeholder_run = p_placeholder.add_run("[CURVA DE EXHALACIÓN]")
                placeholder_run.italic = True
        else:
            # Placeholder si no hay curva
            p_placeholder = doc.add_paragraph()
            p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder_run = p_placeholder.add_run("[CURVA DE EXHALACIÓN - NO DISPONIBLE]")
            placeholder_run.italic = True
        
        doc.add_paragraph()  # Espacio
        
        # ==================== REFERENCIAS ====================
        # Línea separadora
        doc.add_paragraph("_" * 80)
        
        # Texto de referencias
        ref_text = """Interpretation of Exhaled Nitric Oxide Levels (FeNO) for Clinical Applications, Am J Crit CareMed: Vol 184.Pp 602-615, 2011"""
        
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run(ref_text)
        ref_run.font.name = 'Calibri'
        ref_run.font.size = Pt(9)
        ref_run.italic = True
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ==================== GUARDAR DOCUMENTO ====================
        # Guardar en archivo temporal
        if output_path:
            doc.save(output_path)
            return output_path
        else:
            # Devolver bytes
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            with open(temp_file.name, 'rb') as f:
                doc_bytes = f.read()
            
            os.unlink(temp_file.name)
            return doc_bytes
            
    except Exception as e:
        st.error(f"Error generando documento: {str(e)}")
        return None

# ==========================================================
# 4. INTERFAZ STREAMLIT MEJORADA
# ==========================================================
def main():
    # Inicializar variables de sesión
    if 'datos_pdf' not in st.session_state:
        st.session_state.datos_pdf = None
    if 'debug_mode' not in st.session_state:
        st.session_state.debug_mode = False
    
    # Encabezado principal
    st.markdown("<h1 class='main-header'>🏥 INT – Laboratorio de Función Pulmonar</h1>", unsafe_allow_html=True)
    st.markdown("<h3 class='main-header'>Generador de Informes de Óxido Nítrico Exhalado (FeNO)</h3>", unsafe_allow_html=True)
    
    # Opción de debug (oculta)
    with st.sidebar:
        if st.checkbox("Modo Debug", key="debug_checkbox"):
            st.session_state.debug_mode = True
    
    # Crear pestañas para mejor organización
    tab1, tab2, tab3 = st.tabs(["📋 Datos del Paciente", "📄 Carga de PDF", "📊 Vista Previa"])
    
    with tab1:
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("👤 Información Personal")
            nombre = st.text_input("Nombre *", key="nombre")
            apellidos = st.text_input("Apellidos *", key="apellidos")
            rut = st.text_input("RUT/Pasaporte *", key="rut")
            genero = st.selectbox("Género *", ["Seleccione", "Hombre", "Mujer"], key="genero")
            procedencia = st.text_input("Procedencia *", key="procedencia", value="Poli")
        
        with col2:
            st.subheader("📈 Datos Clínicos")
            f_nacimiento = st.text_input("Fecha de Nacimiento (DD/MM/AAAA) *", key="f_nac")
            edad = st.text_input("Edad (años)", key="edad")
            altura = st.text_input("Altura (cm) *", key="altura")
            peso = st.text_input("Peso (kg) *", key="peso")
            medico = st.text_input("Médico Solicitante *", key="medico")
            operador = st.text_input("Operador *", value="TM Jorge Espinoza", key="operador")
            fecha_examen = st.date_input("Fecha del Examen *", key="fecha_examen")
    
    with tab2:
        st.subheader("📤 Cargar Informe Sunvou")
        
        col_upload, col_info = st.columns([2, 1])
        
        with col_upload:
            pdf_file = st.file_uploader(
                "Seleccionar archivo PDF",
                type=["pdf"],
                help="Suba el informe PDF generado por el equipo Sunvou",
                key="pdf_uploader"
            )
            
            if pdf_file:
                st.success(f"✅ Archivo cargado: {pdf_file.name}")
                
                # Mostrar información del PDF
                with col_info:
                    st.markdown("<div class='metric-box'>", unsafe_allow_html=True)
                    st.metric("Tamaño", f"{pdf_file.size / 1024:.1f} KB")
                    st.markdown("</div>", unsafe_allow_html=True)
                
                # Botón para extraer datos
                if st.button("🔍 Extraer Datos del PDF", type="secondary", key="extraer_btn"):
                    with st.spinner("Procesando PDF..."):
                        datos_extractos = extraer_datos_pdf_sunvou(pdf_file)
                        
                        if datos_extractos:
                            st.session_state.datos_pdf = datos_extractos
                            
                            # Mostrar datos extraídos
                            st.markdown("<div class='success-box'>", unsafe_allow_html=True)
                            st.success("✅ Datos extraídos correctamente")
                            
                            col_val1, col_val2, col_val3, col_val4 = st.columns(4)
                            with col_val1:
                                st.metric("FeNO50", f"{datos_extractos['FeNO50']} ppb")
                            with col_val2:
                                st.metric("Temperatura", f"{datos_extractos['Temperatura']} °C")
                            with col_val3:
                                st.metric("Presión", f"{datos_extractos['Presion']} cmH2O")
                            with col_val4:
                                st.metric("Flujo", f"{datos_extractos['Flujo']} ml/s")
                            
                            st.markdown("</div>", unsafe_allow_html=True)
        
        # Opciones de procesamiento
        tipo_prueba = st.radio(
            "Tipo de Prueba",
            ["FeNO 50", "FeNO 50-200"],
            horizontal=True,
            key="tipo_prueba"
        )
    
    with tab3:
        st.subheader("👁️ Vista Previa de Datos")
        
        # Mostrar datos del paciente
        if nombre or rut:
            st.markdown("**📋 Resumen de Datos del Paciente:**")
            col_res1, col_res2 = st.columns(2)
            
            with col_res1:
                st.write(f"**Nombre:** {nombre if nombre else '---'} {apellidos if apellidos else ''}")
                st.write(f"**RUT:** {rut if rut else '---'}")
                st.write(f"**Género:** {genero if genero != 'Seleccione' else '---'}")
                st.write(f"**Edad:** {edad if edad else '---'} años")
            
            with col_res2:
                st.write(f"**Altura:** {altura if altura else '---'} cm")
                st.write(f"**Peso:** {peso if peso else '---'} kg")
                st.write(f"**Procedencia:** {procedencia if procedencia else '---'}")
                st.write(f"**Médico:** {medico if medico else '---'}")
        
        # Mostrar datos extraídos del PDF
        if st.session_state.datos_pdf:
            st.markdown("**📄 Datos Extraídos del PDF:**")
            
            col_data1, col_data2, col_data3, col_data4 = st.columns(4)
            
            with col_data1:
                st.info(f"**FeNO50:** {st.session_state.datos_pdf['FeNO50']} ppb")
            with col_data2:
                st.info(f"**Temperatura:** {st.session_state.datos_pdf['Temperatura']} °C")
            with col_data3:
                st.info(f"**Presión:** {st.session_state.datos_pdf['Presion']} cmH2O")
            with col_data4:
                st.info(f"**Flujo:** {st.session_state.datos_pdf['Flujo']} ml/s")
            
            # Mostrar imágenes si están disponibles
            if st.session_state.datos_pdf.get('img_curva'):
                st.markdown("**📈 Curva de Exhalación:**")
                try:
                    st.image(st.session_state.datos_pdf['img_curva'], caption="Curva extraída del PDF", width=400)
                except:
                    st.warning("No se pudo mostrar la curva de exhalación")
    
    # ==================== BOTÓN DE GENERACIÓN ====================
    st.markdown("---")
    
    col_gen1, col_gen2, col_gen3 = st.columns([1, 2, 1])
    
    with col_gen2:
        if st.button("🚀 GENERAR INFORME COMPLETO", type="primary", use_container_width=True):
            # Validación de datos obligatorios
            campos_obligatorios = {
                "Nombre": nombre,
                "Apellidos": apellidos,
                "RUT": rut,
                "Género": genero,
                "Médico": medico,
                "Operador": operador
            }
            
            campos_faltantes = [k for k, v in campos_obligatorios.items() if not v or (k == "Género" and v == "Seleccione")]
            
            if campos_faltantes:
                st.error(f"❌ Faltan campos obligatorios: {', '.join(campos_faltantes)}")
            elif not pdf_file:
                st.error("❌ Debe cargar un archivo PDF")
            elif not st.session_state.datos_pdf:
                st.error("❌ Debe extraer los datos del PDF primero (haga clic en 'Extraer Datos del PDF')")
            else:
                with st.spinner("🔄 Generando informe profesional..."):
                    try:
                        # Preparar datos completos
                        datos_completos = {
                            "nombre": nombre,
                            "apellidos": apellidos,
                            "rut": rut,
                            "genero": genero,
                            "procedencia": procedencia,
                            "f_nacimiento": f_nacimiento,
                            "edad": edad,
                            "altura": altura,
                            "peso": peso,
                            "medico": medico,
                            "operador": operador,
                            "fecha_examen": fecha_examen.strftime("%d/%m/%Y") if fecha_examen else "---",
                            "feno50": st.session_state.datos_pdf.get("FeNO50", "---"),
                            "temperatura": st.session_state.datos_pdf.get("Temperatura", "---"),
                            "presion": st.session_state.datos_pdf.get("Presion", "---"),
                            "flujo": st.session_state.datos_pdf.get("Flujo", "---"),
                            "img_logo": st.session_state.datos_pdf.get("img_logo"),
                            "img_curva": st.session_state.datos_pdf.get("img_curva")
                        }
                        
                        # Generar documento
                        doc_bytes = generar_documento_identico(datos_completos, output_path=None)
                        
                        if doc_bytes:
                            st.markdown("<div class='success-box'>", unsafe_allow_html=True)
                            st.success("🎉 ¡Informe generado exitosamente!")
                            
                            # Botón de descarga
                            st.download_button(
                                label="📥 DESCARGAR INFORME EN WORD",
                                data=doc_bytes,
                                file_name=f"INFORME_FENO_{rut if rut else 'SIN_RUT'}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary",
                                use_container_width=True
                            )
                            
                            st.info("💡 **Siguientes pasos:** Abra el documento en Word y guárdelo como PDF para distribución.")
                            st.markdown("</div>", unsafe_allow_html=True)
                            
                    except Exception as e:
                        st.error(f"❌ Error durante la generación: {str(e)}")
                        if st.session_state.debug_mode:
                            st.exception(e)
    
    # ==================== INSTRUCCIONES ====================
    with st.expander("📖 Instrucciones de Uso"):
        st.markdown("""
        1. **Complete todos los campos** en la pestaña "Datos del Paciente" (los marcados con * son obligatorios)
        2. **Suba el PDF** del informe Sunvou en la pestaña "Carga de PDF"
        3. **Haga clic en "Extraer Datos del PDF"** para procesar la información
        4. **Verifique los datos** en la pestaña "Vista Previa"
        5. **Haga clic en "GENERAR INFORME COMPLETO"** para crear el documento Word
        6. **Descargue el archivo** y ábralo en Microsoft Word
        7. **Guarde como PDF** (Archivo → Guardar como → PDF) para distribución
        
        **Nota:** El sistema mantiene el formato profesional exacto del modelo institucional.
        """)
    
    # ==================== PIE DE PÁGINA ====================
    st.markdown("---")
    st.caption("© INT - Laboratorio de Función Pulmonar | Sistema de Generación de Informes FeNO v1.0")

# ==========================================================
# 5. FUNCIÓN PRINCIPAL
# ==========================================================
if __name__ == "__main__":
    main()
